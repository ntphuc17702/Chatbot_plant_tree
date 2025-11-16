from docx import Document
import os
import re
BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # thư mục chứa file hiện tại
DATA_DIR = os.path.join(BASE_DIR, "..", "data_output")
def split_docx_by_content(input_path, output_dir=DATA_DIR):
    """
    Tách file .docx theo logic:
      - Heading có text → tách file, giữ heading cha
      - Heading không có text → không tách, giữ trong file gốc
      - File chỉ toàn heading → lưu nguyên file
      - File không có heading → lưu nguyên file
    """
    os.makedirs(output_dir, exist_ok=True)

    parent_folder = os.path.basename(os.path.dirname(input_path))
    file_name = os.path.basename(input_path)
    doc = Document(input_path)

    parts = []
    current_part = []
    current_title = None
    has_heading = False
    has_text_anywhere = False

    def is_text(p):
        return not (p.style.name.startswith("Heading")) and p.text.strip()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if style_name.startswith("Heading"):
            has_heading = True
            # Nếu phần trước có text → lưu
            if current_part and any(is_text(p) for p in current_part):
                parts.append((current_title, current_part))
                current_part = []
            current_title = re.sub(r'[\\/*?:"<>|]', "_", text) or f"part_{len(parts)+1}"
            current_part.append(para)
        else:
            if text:
                has_text_anywhere = True
            current_part.append(para)

    # Thêm phần cuối nếu có text
    if current_part and any(is_text(p) for p in current_part):
        parts.append((current_title, current_part))

    # Nếu file chỉ toàn heading → lưu nguyên file
    if has_heading and not has_text_anywhere:
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)
        print(f"📄 File chỉ toàn heading → đã lưu nguyên: {output_path}")
        return

    # Nếu file không có heading → lưu nguyên file
    if not has_heading:
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)
        print(f"📄 File không có heading → đã lưu nguyên: {output_path}")
        return

    # Ghi các phần có text
    for i, (title, paras) in enumerate(parts, start=1):
        new_doc = Document()
        info = new_doc.add_paragraph(f"{parent_folder} {file_name}")
        info.runs[0].bold = True
        new_doc.add_paragraph("")
        for p in paras:
            new_p = new_doc.add_paragraph(p.text)
            new_p.style = p.style

        safe_name = re.sub(r'[\\/*?:"<>|]', "_", title or f"part_{i}")
        output_path = os.path.join(output_dir, f"{i:02d}_{safe_name}.docx")
        new_doc.save(output_path)
        print(f"✅ Đã tạo: {output_path}")

    print(f"🎯 Hoàn tất: {len(parts)} phần có nội dung trong {file_name}\n")
